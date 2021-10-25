from typing import Text
from docx.api import _default_docx_path
from numpy import NaN, nan
from docx import Document
from funciones.estadistica import *
#from estadistica import *
import docx
import pandas as pd
from docx.shared import Cm
from PyQt5.QtWidgets import QFileDialog

def fill_fields(archivo, url_img1,url_img2,text):
    # create the table 
    table = archivo.add_table(2,2)
    
    # merge second table
    A = table.cell(1,0).merge(table.cell(1,1))
    
    # add first image
    paragraph = table.cell(0,0).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(url_img1,height = Cm(5.4) )
    
    # add second image
    paragraph = table.cell(0,1).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(url_img2,height = Cm(5.4)  )
    
    # add descripction
    A.add_paragraph(text)
    archivo.add_paragraph()
    return archivo

def llenar_info_general():
    df = pd.read_csv('./archivos/current_general.csv',sep= ';',encoding= "latin")
    name = str(df["numero_campo"][0])
    nombre_archivo = QFileDialog.getSaveFileName(directory= name,filter= "documents (*.docx *.doc)")[0]
    parametros = df.iloc[0].tolist()
    parametros.pop()
    parametros.pop()
    parametros.insert(11,"")
    parametros = list(map(str, parametros))
    # archivo = Document("./archivos/templates/template_1.docx") PREGUNTARLE A JOHAN
    archivo = Document()
    archivo.add_heading('INFORMACIÓN GENERAL' )
    archivo.add_paragraph()
    campos = ["IGM", "Número de campo", "Unidad litoestratigráfica", "Localidad",
              "Departamento", "Municipio", "Plancha", "Escala", "Coordenada X", 
              "Origen de Coordenadas", "Coordenada Y" ,"", "Colector", "Fecha de recolección de la muestra",
              "Analizador", "Fecha de Análisis petrográfico", "Número de puntos de conteo"]

    tabla_info_g = archivo.add_table(9,2)
    contador = 0
    for i in range(0,len(campos)//2):
        for j in range(2):
            if campos[contador] == "":
                tabla_info_g.cell(i,j%2).paragraphs[0].add_run(campos[contador]+" "+parametros[contador])
            else:
                p1 = tabla_info_g.cell(i,j%2).paragraphs[0]
                r1 = p1.add_run(campos[contador]+": ")
                r1.bold = True
                p1.add_run(parametros[contador])
            contador +=1
    tabla_info_g.cell(8,0).paragraphs[0].add_run(campos[len(campos)-1]+" "+parametros[-1])
    archivo.add_paragraph()
    archivo.save(nombre_archivo)
    return nombre_archivo

def llenar_macro(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep = ";", encoding= "latin")
    if general["Tipo_r"][0] != "Sedimentaria":
        df = pd.read_csv("./archivos/current_macro.csv", sep = ";", encoding= "latin")
    else:
        df = pd.read_csv("./archivos/current_macro_sed.csv", sep = ";", encoding= "latin")
    campos = df.columns.tolist()
    parametros = df.values[-1].tolist()
    parametros = list(map(str, parametros))
    img_macro = df["url_foto"][0]
    escala = df["url_escala"]
    for i in range(2):
        campos.pop()
        parametros.pop()
    for i in range(len(campos)):
        campos[i] = campos[i].replace("_", " ")
    archivo = Document(nombre_archivo)
    archivo.add_heading('DESCRIPCIÓN MACROSCÓPICA' )
    if len(campos) > 2:
        tabla_macro = archivo.add_table(len(campos),3)
        imagen = tabla_macro.cell(0,2).merge(tabla_macro.cell(len(campos)-1,2))
        parag = imagen.paragraphs[0]
        run = parag.add_run()
        run.add_picture(img_macro, height =Cm(4))
        for i in range(0,len(parametros)):
            for j in range(2):
                if j == 0:
                    tabla_macro.cell(i,j).paragraphs[0].add_run(campos[i])
                if j ==1:
                    tabla_macro.cell(i,j).paragraphs[0].add_run(parametros[i])
    else:
        tabla_macro = archivo.add_table(1,2)
        imagen = tabla_macro.cell(0,1)
        parag = imagen.paragraphs[0]
        run = parag.add_run()
        run.add_picture(img_macro, height = Cm(4))
        tabla_macro.cell(0,0).paragraphs[0].add_run(parametros[0])

    archivo.save(nombre_archivo)


def llenar_inter_plut(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo = Document(nombre_archivo)
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA")
    archivo.add_paragraph()
    subs = ["Textura general:", "Otras texturas o texturas especiales:", 
            "Descripción de la matriz:"]
    for i in subs:
        p1 = archivo.add_paragraph()
        r1 = p1.add_run(i)
        r1.bold = True
        r1.underline = True
        if i == subs[0]:
            p1.add_run(" Mencionar ordenada, clara y correctamente los términos texturales que indican" 
                        + "dimensión relativa de los cristales, tamaños de cristales, cristalinidad, "
                        + "relaciones mutuas entre cristales,  formas cristalinas o desarrollo relativo de caras de cristales")
        else:
            p1.add_run(" ")
        archivo.add_paragraph()
    archivo.add_page_break()
    archivo.add_paragraph()
    archivo.add_heading("COMPOSICIÓN MINERALÓGICA (% VOL) - " + igm)
    archivo.add_paragraph()
    rows = 12
    tabla_perc = archivo.add_table(rows,6)
    tabla_perc.style = 'TableGrid'
    for i in range(rows):
        for j in range(6):
            if j % 2 != 0:
                tabla_perc.cell(i,j).width = Cm(1)
            else:
                tabla_perc.cell(i,j).width = Cm(4)
    for i in range(6):
        if i%2 != 0:
            colum = tabla_perc.columns[i]
            colum.width = Cm(1)
    cells = [0,1,2,3,4,4,0,2]
    content = ["MINERALES \nPRINCIPALES", "%", "MINERALES \nACCESORIOS", "%",
                "MINERALES DE ALTERACIÓN", "MINERALES DE INTRODUCCIÓN", "TOTAL", "TOTAL"]
    
    for i in range(len(cells)):
        if i >5:
            para = tabla_perc.cell(11,cells[i]).paragraphs[0]
        elif i == 5:
            para = tabla_perc.cell(6,cells[i]).paragraphs[0]
        else:
            para = tabla_perc.cell(0,cells[i]).paragraphs[0]
        r1 = para.add_run(content[i])
        r1.bold = True
    tabla_perc.cell(0,4).merge(tabla_perc.cell(0,5))
    tabla_perc.cell(6,4).merge(tabla_perc.cell(6,5))
       
    archivo.add_paragraph()
    clas_roc = archivo.add_paragraph()
    run = clas_roc.add_run("CLASIFICACIÓN DE LA ROCA ")
    run.bold = True
    run.underline = True
    r2 = clas_roc.add_run("(Basada en Streckeisen, 1976):")
    r2.underline = True
    clas_roc.add_run(" La clasificación")
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA DE MINERALES")
    archivo.add_paragraph()
    p1 = archivo.add_paragraph()
    r1 = p1.add_run("Mineral 1:")
    r1.bold = True
    r1.underline = True
    p1.add_run(" Descripción concisa y completa de rasgos generales y particulares," 
                + "sin olvidar tamaño, forma, color, distribución, relaciones texturales, "
                + "extinción, clivaje, etc")
    archivo.add_paragraph()
    archivo.add_heading("OBSERVACIONES")
    archivo.add_paragraph()
    archivo.save(nombre_archivo)

def llenar_inter_volcanico(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo = Document(nombre_archivo)
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA")
    archivo.add_paragraph()
    lista=['Textura general:','Otras texturas o Texturas especiales:','Descripción de la matriz:']
    for i in lista:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.bold = True
        r1.underline= True
        p1.add_run(" ")  
    archivo.add_paragraph()
    archivo.add_heading("COMPOSICIÓN MINERALÓGICA (% Vol.) - " + igm)
    archivo.add_paragraph()
    tabla_comp_volc= archivo.add_table(14,5)
    tabla_comp_volc.style = 'TableGrid'
    tabla_comp_volc.cell(1,0).merge(tabla_comp_volc.cell(1,4))
    tabla_comp_volc.cell(5,0).merge(tabla_comp_volc.cell(5,4))
    tabla_comp_volc.cell(8,0).merge(tabla_comp_volc.cell(8,4))
    tabla_comp_volc.cell(11,0).merge(tabla_comp_volc.cell(11,1))
    tabla_comp_volc.cell(11,2).merge(tabla_comp_volc.cell(11,4))
    tabla_comp_volc.cell(12,2).merge(tabla_comp_volc.cell(13,4))
    lista=['Componentes', 'Total en roca \n(% )','Fenocristal\n( ≥ 2 mm)',
          'Microfenocristal\n(2 – 0,5mm)','Microcristal\n( < 0,5mm)']
    for i in range (len(lista)):
        if i==i:
            a=tabla_comp_volc.cell(0,i).paragraphs[0] 
            r1= a.add_run(lista [i])
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
    a=tabla_comp_volc.cell(1,0).paragraphs[0] 
    r1= a.add_run('Minerales principales')
    r1.bold = True  
    a=tabla_comp_volc.cell(5,0).paragraphs[0] 
    r1= a.add_run('Minerales accesorios')
    r1.bold = True
    a=tabla_comp_volc.cell(8,0).paragraphs[0] 
    r1= a.add_run('Minerales de alteración')
    r1.bold = True
    a=tabla_comp_volc.cell(11,0).paragraphs[0] 
    r1= a.add_run('Matriz:')
    r1.bold = True
    a=tabla_comp_volc.cell(11,2).paragraphs[0] 
    r1= a.add_run('Observaciones:')
    r1.bold = True
    tabla_comp_volc.cell(12,0).text=('Fracción Criptocristalina')
    tabla_comp_volc.cell(13,0).text=('Vidrio')
    archivo.add_paragraph()
    tabla_perc_volc= archivo.add_table(1,6)
    tabla_perc_volc.style = 'TableGrid'
    perc_vol= ['(%) de cristales:','(%) de matriz:','(%) de vesículas:']
    contador=0
    for i in range (6):
        if i%2==0:
            a=tabla_perc_volc.cell(0,i).paragraphs[0] 
            r1= a.add_run(perc_vol[contador])
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
            contador+=1
    archivo.add_paragraph()
    clas_roc = archivo.add_paragraph()
    run = clas_roc.add_run("CLASIFICACIÓN DE LA ROCA ")
    run.bold = True
    run.underline = True
    r2 = clas_roc.add_run("(Basada en Streckeisen, 1978):")
    r2.underline = True
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA DE MINERALES")
    archivo.add_paragraph()
    p1 = archivo.add_paragraph()
    r1 = p1.add_run("Mineral 1:")
    r1.bold = True
    r1.underline = True
    p1.add_run(" Descripción concisa y completa de rasgos generales y particulares," 
                + "sin olvidar tamaño, forma, color, distribución, relaciones texturales, "
                + "extinción, clivaje, etc")
    archivo.add_paragraph()
    archivo.add_heading("OBSERVACIONES")
    archivo.add_paragraph()
    
    archivo.save(nombre_archivo)

def llenar_inter_volcclas(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo=Document(nombre_archivo)
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA")
    archivo.add_paragraph()
    lista=['Textura general:','Otras texturas:','Descripción de la matriz:', 'Observaciones:']
    for i in lista:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.bold = True
        r1.underline= True
        p1.add_run(" ")  
    archivo.add_paragraph()
    archivo.add_heading("COMPOSICIÓN (% Vol.) - " + igm)
    archivo.add_paragraph()
    tabla_comp_volcclas= archivo.add_table(19,3)
    tabla_comp_volcclas.style = 'TableGrid'
    for i in range (8):
        tabla_comp_volcclas.cell(i,0).merge(tabla_comp_volcclas.cell(i,1))
    tabla_comp_volcclas.cell(8,0).merge(tabla_comp_volcclas.cell(16,0))
    tabla_comp_volcclas.cell(1,0).merge(tabla_comp_volcclas.cell(1,2))
    tabla_comp_volcclas.cell(5,0).merge(tabla_comp_volcclas.cell(5,2))
    tabla_comp_volcclas.cell(17,0).merge(tabla_comp_volcclas.cell(17,2))
    tabla_comp_volcclas.cell(18,0).merge(tabla_comp_volcclas.cell(18,2))
    tabla_comp_volcclas.cell(8,1).merge(tabla_comp_volcclas.cell(8,2))
    tabla_comp_volcclas.cell(11,1).merge(tabla_comp_volcclas.cell(11,2))
    tabla_comp_volcclas.cell(14,1).merge(tabla_comp_volcclas.cell(14,2))
    a=tabla_comp_volcclas.cell(0,0).paragraphs[0] 
    r1= a.add_run('Componentes')
    r1.bold = True
    a=tabla_comp_volcclas.cell(0,2).paragraphs[0] 
    r1= a.add_run('Total en roca (%)')
    r1.bold = True
    mid= ['Juveniles','Accesorios','Accidentales']
    contador=0
    for i in range (8,15):
        if (i-8)%3==0:
            a=tabla_comp_volcclas.cell(i,2).paragraphs[0] 
            r1= a.add_run(mid[contador])
            r1.bold = True
            contador+=1
    bot= ['Cristales y Fragmentos cristalinos','Fragmentos vesiculados y vidrio','Fragmentos líticos    ']
    contador=0
    for i in range (1,12):
        if (i-1)%4==0:
            a=tabla_comp_volcclas.cell(i,0).paragraphs[0] 
            r1= a.add_run(bot[contador])
            r1.bold = True
            contador+=1
    a=tabla_comp_volcclas.cell(17,0).paragraphs[0] 
    r1= a.add_run('Observaciones')
    r1.bold = True
    archivo.add_heading('Nota: Este formato se usará para la lámina delgada de la muestra de una ignimbrita' 
    'o de la matriz de un depósito volcaniclástico. Para el caso de un fragmento lítico o pumítico,' 
    'tomado de un depósito piroclástico, se usaría el formato de roca volcánica',4)
    archivo.add_paragraph()
    archivo.add_paragraph()
    clas_roc = archivo.add_paragraph()
    run = clas_roc.add_run("CLASIFICACIÓN DE LA MUESTRA ")
    run.bold = True
    run.underline = True
    r2 = clas_roc.add_run("(Según criterios de campo, macroscópicos y microscópicos):")
    r2.underline = True
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA LOS COMPONENTES")
    archivo.add_paragraph()
    p1 = archivo.add_paragraph()
    r1 = p1.add_run("Nombre componente 1:")
    r1.bold = True
    r1.underline = True
    r1.italic= True
    p1.add_run(" Descripción concisa y completa de rasgos generales y particulares," 
                + "sin olvidar tamaño, forma, color, distribución, relaciones texturales, "
                + "extinción, clivaje, etc")
    archivo.add_paragraph()
    archivo.add_heading("OBSERVACIONES:")
    archivo.add_heading('(Indicar rasgos particulares y especialmente aquellos que indiquen'
                        'eventos que afectaron a la roca con posterioridad a su formación por ejemplo: '
                        'intrusión de venillas, alteraciones hidrotermales, microfallamiento, recristalización '
                        'térmica, deformación, etc. – Si ellos se dan en la muestra.) ',4)
    archivo.add_paragraph()
    archivo.save(nombre_archivo)


def llenar_inter_dinamico(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo = Document(nombre_archivo)
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA")
    archivo.add_paragraph()
    archivo.add_heading('ASPECTOS TEXTURALES',2)
    archivo.add_paragraph()
    archivo.add_heading('DESCRIPCIÓN DE MICROESTRUCTURAS: ',3)
    archivo.add_paragraph()
    archivo.add_heading('FABRICA: ',3)
    list=['Cohesión _____','Foliación_____','Tamaño de porfiroclastos _____ (µ)'
    'Proporción de matriz____%','Recristalización_____']
    for i in list:
        p1= archivo.add_paragraph(i)
        p1.add_run(" ")
    archivo.add_paragraph()
    archivo.add_heading('TIPO DE DEFORMACIÓN:',3)
    archivo.add_paragraph('Frágil _____')
    archivo.add_paragraph('Dúctil_____  Grado: Bajo	_____Medio_____Alto_____')
    archivo.add_paragraph()
    archivo.add_heading('ASPECTOS COMPOSICIONALES - ' + igm,2)
    archivo.add_paragraph()
    tabla_comp_din= archivo.add_table(10,10)
    tabla_comp_din.cell(0,0).merge(tabla_comp_din.cell(1,0))
    tabla_comp_din.cell(0,1).merge(tabla_comp_din.cell(1,1))
    tabla_comp_din.cell(0,2).merge(tabla_comp_din.cell(0,5))
    tabla_comp_din.cell(0,6).merge(tabla_comp_din.cell(0,9))
    tabla_comp_din.style = 'TableGrid'
    content = ["PORFIROCLASTOS", "%", "COMPOSICIÓN MATRIZ"] 
    content2= ["Relíctica", "%", "Neo-formada", "%","De Alteración", "%",'De Introducción', "%"]
    for i in range (len(content)):
        if i==i:
            a=tabla_comp_din.cell(0,i).paragraphs[0] 
            r1= a.add_run(content [i])
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
    
    a=tabla_comp_din.cell(0,6).paragraphs[0] 
    r1= a.add_run('MINERALES SECUNDARIOS')
    r1.italic = True
    r1.bold = True   
    
    for i in range (len(content2)):
        if i==i:
            a=tabla_comp_din.cell(1,i+2).paragraphs[0] 
            r1= a.add_run(content2 [i])
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
    for i in range (10):
        if i%2==0:
            a=tabla_comp_din.cell(9,i).paragraphs[0] 
            r1= a.add_run('TOTAL')
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
    archivo.add_paragraph()
    archivo.add_heading('CLASIFICACIÓN DE LA ROCA',2)
    archivo.add_paragraph()
    autores= ['Higgins (1971):', 'Sibson (1977):', 'Wise & otros (1984):', 
              'Marshak & Mitra (1988):', 'Trouw & otros (2010):']
    for i in autores:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.italic = True
        r1.bold = True
        p1.add_run(" ")    
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA DE MINERALES")
    archivo.add_paragraph()
    p1 = archivo.add_paragraph()
    r1 = p1.add_run("Mineral 1:")
    r1.bold = True
    r1.underline = True
    p1.add_run(" Descripción concisa y completa de rasgos generales y particulares," 
                + "sin olvidar tamaño, forma, color, distribución, relaciones texturales, "
                + "extinción, clivaje, etc")
    archivo.add_paragraph()   

    archivo.save(nombre_archivo)  
    
def llenar_inter_regional(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo = Document(nombre_archivo)
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA")
    archivo.add_paragraph()
    subs = ["Textura general:", "Otras texturas:", "Observaciones:"]
    for i in subs:
        p1 = archivo.add_paragraph()
        r1 = p1.add_run(i)
        r1.bold = True
        r1.underline = True
        p1.add_run(" ")
        archivo.add_paragraph()

    archivo.add_page_break()
    archivo.add_paragraph()
    archivo.add_heading("COMPOSICIÓN MINERALÓGICA (%Vol) - " + igm)
    archivo.add_paragraph()
    rows = 12
    tabla_perc = archivo.add_table(rows,6)
    for i in range(rows):
        for j in range(6):
            if j % 2 != 0:
                tabla_perc.cell(i,j).width = Cm(1)
            else:
                tabla_perc.cell(i,j).width = Cm(4)
    for i in range(6):
        if i%2 != 0:
            colum = tabla_perc.columns[i]
            colum.width = Cm(1)
    tabla_perc.style =  'TableGrid'
    cells = [0,1,2,3,4,4,0,2]
    content = ["MINERALES \nPRINCIPALES", "%", "MINERALES \nACCESORIOS", "%",
                "MINERALES DE ALTERACIÓN", "MINERALES DE INTRODUCCIÓN", "TOTAL", "TOTAL"]
    
    for i in range(len(cells)):
        if i >5:
            para = tabla_perc.cell(11,cells[i]).paragraphs[0]
        elif i == 5:
            para = tabla_perc.cell(6,cells[i]).paragraphs[0]
        else:
            para = tabla_perc.cell(0,cells[i]).paragraphs[0]
        r1 = para.add_run(content[i])
        r1.bold = True
        
    tabla_perc.cell(0,4).merge(tabla_perc.cell(0,5))
    tabla_perc.cell(6,4).merge(tabla_perc.cell(6,5))
    archivo.add_paragraph()
    caracteristicas = ["PARAGÉNESIS:", "TIPO DE METAMORFÍSMO:",
                       "FACIES DE METAMORFISMO:", "PROTOLITO:"]
    for i in caracteristicas:
        p1 = archivo.add_paragraph()
        r1 = p1.add_run(i)
        r1.bold = True
        r1.underline = True
        p1.add_run(" ")
    clas_roc = archivo.add_paragraph()
    run = clas_roc.add_run("CLASIFICACIÓN DE LA ROCA ")
    run.bold = True
    run.underline = True
    r2 = clas_roc.add_run("(Basada en SSMR, 2007):")
    r2.underline = True
    clas_roc.add_run(" La clasificación")
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA DE MINERALES")
    archivo.add_paragraph()
    p1 = archivo.add_paragraph()
    r1 = p1.add_run("Mineral 1:")
    r1.bold = True
    r1.underline = True
    p1.add_run(" Descripción concisa y completa de rasgos generales y particulares, "
                + "sin olvidar tamaño, forma, color, distribución, relaciones texturales, "
                + "extinción, clivaje, etc")
    archivo.add_paragraph()
    archivo.save(nombre_archivo)

def llenar_inter_silici(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    campos_form = ["Metamorficos", "Volcanicos", "Plutonicos", "Sedimentarios", "Primaria", "Secundaria",
                   "Cuarzo mono", "Cuarzo poli", "Chert", "Feldespato K", "Feldespato Na - Ca", "Micas",
                   "Min. arcillosos", "Granos aloq.", "Otros terrigenos", "Opacos", "Materia org.",
                   "Cemento", "Otros ortoq.", "Porosidad", "Cuarzo", "Liticos", "Feldespato","Terrigenos"]
    porosidad = ["Primaria", "Secundaria"]
    feldes = ["Feldespato K", "Feldespato Na - Ca"]
    litic = ["Metamorficos", "Volcanicos", "Plutonicos", "Sedimentarios"]
    quartz = ["Cuarzo mono", "Cuarzo poli"]
    datos_comp = datos_silic()
    complete = ["Porosidad", "Cuarzo", "Liticos", "Feldespato", "Terrigenos"]
    terrig = ["Cuarzo mono", "Cuarzo poli", "Chert", "Feldespato K", "Feldespato Na - Ca", "Micas",
              "Min. arcillosos", "Granos aloq.", "Otros terrigenos", "Opacos"]
    for i in complete:
        datos_comp[i] = 0.00
    for i in porosidad:
        try:
            datos_comp["Porosidad"] += datos_comp[i]
        except:
            continue
    for i in feldes:
        try:
            datos_comp["Feldespato"] += datos_comp[i]
        except:
            continue

    for i in quartz:
        try:
            datos_comp["Cuarzo"] += datos_comp[i]
        except:
            continue
    for i in litic:
        try:
            datos_comp["Liticos"] += datos_comp[i]
        except:
            continue
    for i in terrig:
        try:
            datos_comp["Terrigenos"] += datos_comp[i]
        except:
            continue
    for i in campos_form:
        try:
            datos_comp[i] = str(datos_comp[i])
        except:
            datos_comp[i] = str(0.00)

    archivo = Document(nombre_archivo)
    archivo.add_page_break()
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA - "  + igm)
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN TEXTURAL",2)
    archivo.add_paragraph()
    lista= ['HOMOGENEIDAD DE LA ROCA:','TAMAÑO DE GRANO PROMEDIO:',
            'RANGO DE TAMAÑOS:','SELECCIÓN:','REDONDEZ PROMEDIO:',
            'ESFERICIDAD PROMEDIO:', 'MADUREZ TEXTURAL:']
    for i in lista:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.italic = True
        r1.bold = True
        p1.add_run(" ")
    p1 = archivo.add_paragraph()
    list= ['GRAVA ______ (%)','Tamaño promedio:  ____	 Redondez:  _____ Esfericidad:  ______', 'ARENA _____(%)', 
           'Tamaño promedio:  ____	 Redondez:  _____ Esfericidad:  ______', 'LODO ________(%)',
           'Arcilla__________% Tamaño promedio fracción arcilla:______mm/μm\nLimo __________% Tamaño promedio fracción limo:______mm/μm',
           'CONTACTO ENTRE GRANOS:', 'Flotante:______% Tangencial:______%\nLongitudinal:______% Cóncavo-convexo:______%\nSuturado:______%',
           'SOPORTE DE LA ROCA:', 'Granos terrígenos-aloquímicos_________% Minerales arcillosos_________%' ]
    
    contador = 0
    for i in list: #los datos de grava y tamaño y lo demás de la list se pueden llenar con info de interfaz
        if contador %2 == 0:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
        else:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            p1.add_run(" ")
        archivo.add_paragraph()    
        contador = contador +1 
    p1 = archivo.add_paragraph()
    r1 = p1.add_run("POROSIDAD:") 
    r1.bold = True
    r1.underline = True
    p1.add_run(" " +datos_comp["Porosidad"]+ ' %\tPrimaria: ' + datos_comp["Primaria"] + ' %\tSecundaria: ' +datos_comp["Secundaria"]+ 
    ' % \nTipo(s), origen y descripción')
    archivo.add_paragraph()
    p2 = archivo.add_paragraph()
    r1 = p2.add_run("ESTRUCTURAS:") 
    r1.bold = True
    r1.underline = True
    archivo.add_paragraph()
    archivo.add_heading('CLASIFICACIÓN TEXTURAL')
    archivo.add_paragraph()
    clas_roc = archivo.add_paragraph()
    run = clas_roc.add_run("NOMBRE TEXTURAL ")
    run.bold = True
    run.underline = True
    r2 = clas_roc.add_run("(Folk, 1954):")
    r2.underline = True
    archivo.add_paragraph('(Grava + Arena + Lodo = 100%)')
    archivo.add_page_break()
    archivo.add_paragraph()
    archivo.add_heading("DESCRIPCIÓN COMPOSICIONAL - "  + igm)
    archivo.add_paragraph()
    archivo.add_heading("TERRIGENOS: " + datos_comp["Terrigenos"] + '(%)',3)
    archivo.add_paragraph()

    list2= ['Cuarzo: '+ datos_comp["Cuarzo"]+' (%)','Monocristalino: ' + datos_comp["Cuarzo mono"]+ ' (%) Tamaño promedio:___mm/μm Esfericidad:___Redondez:___'
            '\nPolicristalino: '+ datos_comp["Cuarzo poli"]+ ' %	Tamaño promedio:____mm/μm Esfericidad____Redondez:___'
            '\nObservaciones:____ ','Chert: ' + datos_comp["Chert"] + ' %', 'Tamaño promedio: mm/μm Esfericidad_____ Redondez:_____', 'Feldespato: ' + 
            datos_comp["Feldespato"] +' %',
            'Potásico: ' + datos_comp["Feldespato K"] + ' % Tamaño promedio:_____mm/μm	Esfericidad	_____Redondez:_____'	
            '\nSódico-Cálcico: '+ datos_comp["Feldespato Na - Ca"]+ ' % Tamaño promedio:_____mm/μm Esfericidad_____Redondez:_____']
    contador=0
    for i in list2: #los datos de grava y tamaño y lo demás de la list se pueden llenar con info de interfaz
        if contador %2 == 0:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
        else:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            p1.add_run(" ")
        archivo.add_paragraph()    
        contador = contador +1 
    list3=[ 'Micas: ' + datos_comp["Micas"] + ' %','Descripción ______','Minerales Arcillosos: '+datos_comp["Min. arcillosos"] +
            ' %',' Descripción_____','Granos Aloquímicos: '+datos_comp["Granos aloq."] +' %','Descripción_____', 'Otros Terrígenos: ' +
            datos_comp["Otros terrigenos"]+ ' %','Descripción_____','Opacos: '+ datos_comp["Opacos"] +' %','Descripción_______']       
    contador=0
    for i in list3: #los datos de grava y tamaño y lo demás de la list se pueden llenar con info de interfaz
        if contador %2 == 0:
            p1= archivo.add_paragraph(False)
            r1= p1.add_run(i)
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
        else:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            p1.add_run(" ")
        archivo.add_paragraph()    
        contador = contador +1 
    archivo.add_paragraph()
    archivo.add_heading('LÍTICOS (Ígneos, Metamórficos, Sedimentarios): '+datos_comp["Liticos"] +' (%)',3)
    archivo.add_paragraph()
    list4= ['Líticos Metamórficos: ' + datos_comp["Metamorficos"]+ ' %',	
            'Tamaño promedio:_____mm	Esfericidad:_____Redondez:_____',	
            'Líticos Volcánicos: '+ datos_comp["Volcanicos"]+' %',	
            'Tamaño promedio:_____mm	Esfericidad:_____Redondez:_____',	
            'Líticos Plutónicos: '+ datos_comp["Plutonicos"]+ ' %',	
            'Tamaño promedio:_____mm	Esfericidad:_____Redondez:_____',	
            'Líticos Sedimentarios: ' + datos_comp["Sedimentarios"]+ ' %',
            'Tamaño promedio:_____mm	Esfericidad:_____Redondez:_____'	
            '\nObservaciones:_____', 'Materia Orgánica_____%', 'Tipo(s):_____',
            'Cemento: '+ datos_comp["Cemento"]+ ' %','Tipo(s):''\nTamaño cristalino_____mm', 'Otros Ortoquímicos: ' +
             datos_comp["Otros ortoq."]+ ' %', 'Tipo(s)(incluye minerales autigénicos):_____' '\nTamaño:____mm']
    contador=0
    for i in list4: 
        if contador %2 == 0:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
        else:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            p1.add_run(" ")
        archivo.add_paragraph()    
        contador = contador +1 
    archivo.add_paragraph()
    archivo.add_heading('CLASIFICACIÓN COMPOSICIONAL')
    archivo.add_paragraph()
    clas_roc = archivo.add_paragraph()
    run = clas_roc.add_run("NOMBRE COMPOSICIONAL ")
    run.bold = True
    run.underline = True
    r2 = clas_roc.add_run("(Folk, 1954):")
    r2.underline = True
    archivo.add_paragraph()
    archivo.add_heading('DIAGÉNESIS',2)
    z1=archivo.add_paragraph('Autigénesis:')
    z2=archivo.add_paragraph('Recristalización:')
    z1.bold= True
    z2.bold= True
    archivo.save(nombre_archivo)    

def llenar_inter_calc(nombre_archivo):
    archivo = Document(nombre_archivo)
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo.add_page_break()
    archivo.add_heading("DESCRIPCIÓN MICROSCÓPICA - "  + igm)
    archivo.add_paragraph()
    archivo.add_heading("TEXTURA - COMPOSICIÓN",2)
    archivo.add_paragraph()
    lista= ['HOMOGENEIDAD DE LA ROCA:_____','ALOQUÍMICOS_____(%)']
    for i in lista:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.italic = True
        r1.bold = True
        p1.add_run(" ")
    p1 = archivo.add_paragraph()
    lista2= ['Bioclastos_____(%)','Tipo(s), rango de tamaño, selección y redondez:','Peloides _____(%)',
             'Tamaño y origen:','Ooides_____(%)','Tamaño, forma, tipo y estructura interna:','Intraclastos _____(%)',
             'Tamaño, redondez y selección:','Oncoides _____(%)','Tamaño, forma y estructura interna:',
             'Otros Aloquímicos	_____(%)','Tipo(s), tamaño(s), forma(s) y porcentajes:','Terrígenos		(%)',
             'Tipo(s), tamaño(s) y porcentaje(s):','Minerales Autigénicos _____(%)', 'Tipo(s), tamaño(s),'
              'forma(s) y porcentajes:','Extraclastos (fragmentos de rocas carbonatadas) _____(%)',
              'Tipo(s), tamaño(s), redondez y porcentajes:','ORTOQUÍMICOS _____(%)','Tipo(s), distribución'
              'y porcentaje(s):','CEMENTO ESPARÍTICO: _____(%)','Tamaño de cristales, forma de cristales y '
              'distribución:','SOPORTE DE LA ROCA:','Granos- aloquímicos _____%	Lodo calcáreo o micrita	_____% ',
              'CONTACTO ENTRE GRANOS:','Flotante_____% Tangencial_____% \n Longitudinal_____ %'	
              'Cóncavo-convexo_____ %	Suturado_____%','POROSIDAD:_____%','Primaria:_____%	Secundaria:_____% \n'
              'Tipo(s), origen y descripción:']
    contador = 0
    for i in lista2: 
        if contador %2 == 0:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            r1.italic = True
            r1.bold = True
            p1.add_run(" ")
        else:
            p1= archivo.add_paragraph()
            r1= p1.add_run(i)
            p1.add_run(" ")   
        contador = contador +1 
    archivo.add_paragraph()
    archivo.add_heading('CLASIFICACIÓN DE LA ROCA – ' + igm +":")
    archivo.add_paragraph()
    lista3=['Folk (1962):', 'Dunham (1962):', 'Gama textural de Folk (1962):'] 
    for i in lista3:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.italic = True
        r1.bold = True
        p1.add_run(" ") 
    archivo.add_paragraph()
    archivo.add_heading('DIAGÉNESIS:')
    lista4=['Autigénesis:' , 'Recristalización: ']
    for i in lista4:
        p1= archivo.add_paragraph()
        r1= p1.add_run(i)
        r1.italic = True
        r1.bold = True
        p1.add_run(" ") 
    archivo.save(nombre_archivo)
  
def llenar_fotos_micro(nombre_archivo):
    general = pd.read_csv("./archivos/current_general.csv", sep= ";" , encoding= "latin")
    igm = str(general.iloc[0]["igm"])
    if igm == nan: igm = "IGM"
    archivo = Document(nombre_archivo)
    archivo.add_page_break()
    archivo.add_paragraph()
    archivo.add_heading("REGISTRO FOTOGRÁFICO")
    archivo.add_paragraph()
    micro_tab = pd.read_csv("./archivos/current_micro.csv", sep = ";", encoding= "latin")
    rows_n = micro_tab.count()[0]
    for i in range(rows_n):
        url_img1= micro_tab.iloc[i]["url_ppl"]
        url_img2= micro_tab.iloc[i][ "url_xpl"]
        text = micro_tab.iloc[i][ "descrpcion_micro"] 
        text = "IGM: " + igm + ". " + text      
        archivo = fill_fields(archivo,url_img1,url_img2,text)
    archivo.save(nombre_archivo)


# llenar_inter_calc("./test_tabla.docx")
# archivo = Document()
# tabla_macro = archivo.add_table(10,3)
# imagen = tabla_macro.cell(0,2).merge(tabla_macro.cell(8,2))
# parag = imagen.paragraphs[0]
# run = parag.add_run()
# run.add_picture('archivos\Snap-91_PPL.jpg',width = Inches(1),height =Inches(1))
# archivo.add_paragraph("ñandú")
# archivo.save("prueb.docx") 
# nombre_a = "./archivos/perritos.docx"
# llenar_info_general( nombre_a, "IGM:", "Número de campo", "Unidad litoestratigráfica", "Localidad",
#               "Departamento", "Municipio", "Plancha", "Escala", "Coordenada X:", 
#               "Origen de Coordenadas:", "Coordenada Y:" ,"", "Colector", "Fecha de recolección de la muestra",
#               "Analizador", "Fecha de Análisis petrográfico", "Número de puntos de conteo")
 
