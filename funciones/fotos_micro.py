from docx import Document
from docx.shared import Inches

def fill_fields(archivo, url_img1,url_img2,text):
    # create the table 
    table = archivo.add_table(2,2)
    
    # merge second table
    A = table.cell(1,0).merge(table.cell(1,1))
    
    # add first image
    paragraph = table.cell(0,0).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(url_img1,width = Inches(2.95),height =Inches(2.18) )
    
    # add second image
    paragraph = table.cell(0,1).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(url_img2,width = Inches(2.95) ,height =Inches(2.18) )
    
    # add descripction
    A.add_paragraph(text)
    return archivo

# archivo = Document()

# url_img1= 'archivos\Snap-91_PPL.jpg'
# url_img2= 'archivos\Snap-92.jpg'

# text = 'Habia una vez una iguana tomando cafe '
    
# archivo = fill_fields(archivo,url_img1,url_img2,text)
# archivo.save('archivo_geologico.docx')